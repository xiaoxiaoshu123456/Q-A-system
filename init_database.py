"""Core backend utilities for Milvus RAG pipeline.

- Loads and splits documents into parent/child chunks.
- Generates dense and sparse embeddings using FlagEmbedding BGEM3.
- Ensures Milvus collection schema, writes and queries hybrid results.
- Calls LLM according to config.ini provider settings (OpenAI-compatible or Anthropic).
- Exposes a CLI for ingestion, deletion, and query demonstration.
"""
from __future__ import annotations

import argparse
import configparser
import hashlib
import logging
import os
import sys
from collections import deque
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Iterable, Sequence

import numpy as np
from pymilvus import Collection, CollectionSchema, DataType, FieldSchema, connections, utility


@dataclass(frozen=True)
class RetrievalConfig:
    parent_chunk_size: int
    child_chunk_size: int
    chunk_overlap: int
    retrieval_k: int
    candidate_m: int


@dataclass(frozen=True)
class MilvusConfig:
    host: str
    port: int
    database_name: str
    collection_name: str


class Log:
    def __init__(self, max_lines: int = 2000):
        self._lines: deque[str] = deque(maxlen=max(1, int(max_lines)))

    def append(self, line: str) -> None:
        self._lines.append(str(line))

    def clear(self) -> None:
        self._lines.clear()

    def tail(self, n: int = 200) -> str:
        n = max(0, int(n))
        if n == 0:
            return ""
        return "\n".join(list(self._lines)[-n:])

    def lines(self) -> list[str]:
        return list(self._lines)

    def __str__(self) -> str:
        return "\n".join(self._lines)


def _ensure_log_file_handler(logger: logging.Logger, log_file: Path, level: int) -> None:
    log_file = log_file.resolve()
    log_file.parent.mkdir(parents=True, exist_ok=True)
    for h in logger.handlers:
        if isinstance(h, logging.FileHandler):
            try:
                if Path(h.baseFilename).resolve() == log_file:
                    return
            except Exception:
                continue
    handler = logging.FileHandler(str(log_file), encoding="utf-8")
    handler.setLevel(level)
    handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s", "%Y-%m-%d %H:%M:%S"))
    logger.addHandler(handler)


def get_app_logger(
    name: str = "app",
    level: int = logging.INFO,
    log_file: str | Path = Path(__file__).resolve().parent / "logs" / "app.log",
) -> logging.Logger:
    logger = logging.getLogger(name)
    logger.setLevel(level)
    logger.propagate = False
    _ensure_log_file_handler(logger, Path(log_file), level)
    return logger


class RealtimeLogger:
    def __init__(
        self,
        name: str = "app",
        level: int = logging.INFO,
        max_lines: int = 2000,
        stream: Any = None,
        log: Log | None = None,
        log_file: str | Path = Path(__file__).resolve().parent / "logs" / "app.log",
    ):
        self.log = log or Log(max_lines=max_lines)
        self._listener: Callable[[str], None] | None = None

        logger = get_app_logger(name=name, level=level, log_file=log_file)

        if not any(isinstance(h, _RealtimeLogHandler) for h in logger.handlers):
            handler = _RealtimeLogHandler(self.log, stream=stream or sys.stdout)
            handler.setLevel(level)
            handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s", "%Y-%m-%d %H:%M:%S"))
            logger.addHandler(handler)

        self._logger = logger

    def set_listener(self, listener: Callable[[str], None] | None) -> None:
        self._listener = listener
        for h in self._logger.handlers:
            if isinstance(h, _RealtimeLogHandler):
                h.set_listener(listener)

    def tail(self, n: int = 200) -> str:
        return self.log.tail(n)

    def debug(self, msg: str) -> None:
        self._logger.debug(msg)

    def info(self, msg: str) -> None:
        self._logger.info(msg)

    def warning(self, msg: str) -> None:
        self._logger.warning(msg)

    def error(self, msg: str) -> None:
        self._logger.error(msg)

    def exception(self, msg: str) -> None:
        self._logger.exception(msg)


class _RealtimeLogHandler(logging.Handler):
    def __init__(self, log: Log, stream: Any):
        super().__init__()
        self._log = log
        self._stream = stream
        self._listener: Callable[[str], None] | None = None

    def set_listener(self, listener: Callable[[str], None] | None) -> None:
        self._listener = listener

    def emit(self, record: logging.LogRecord) -> None:
        try:
            msg = self.format(record)
        except Exception:
            msg = record.getMessage()
        try:
            self._log.append(msg)
        except Exception:
            pass
        try:
            self._stream.write(msg + "\n")
            self._stream.flush()
        except Exception:
            pass
        if self._listener is not None:
            try:
                self._listener(msg)
            except Exception:
                pass


class DocumentLoader:
    def read_text_file(self, path: Path) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="replace")

    def load_pdf(self, path: Path) -> str:
        import fitz

        doc = fitz.open(str(path))
        parts: list[str] = []
        for page in doc:
            text = page.get_text("text").strip()
            if text:
                parts.append(text)
        return "\n\n".join(parts).strip()

    def load_docx(self, path: Path) -> str:
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([(c.text or "").strip() for c in row.cells]).strip()
                if row_text and row_text != "|":
                    parts.append(row_text)
        return "\n".join(parts).strip()

    def load(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return self.read_text_file(path)
        if suffix == ".pdf":
            text = self.load_pdf(path)
            if not text:
                raise ValueError(f"PDF 解析结果为空: {path}")
            return text
        if suffix == ".docx":
            text = self.load_docx(path)
            if not text:
                raise ValueError(f"Word(.docx) 解析结果为空: {path}")
            return text
        if suffix == ".doc":
            raise ValueError(f"暂不支持 Word(.doc): {path}，请另存为 .docx")
        raise ValueError(f"不支持的文档类型: {suffix} ({path})")


class ChunkBuilder:
    def split_text(self, text: str, chunk_size: int, overlap: int) -> list[str]:
        if chunk_size <= 0:
            raise ValueError("chunk_size 必须大于 0")
        if overlap < 0:
            raise ValueError("overlap 不能小于 0")
        if overlap >= chunk_size:
            raise ValueError("overlap 必须小于 chunk_size")

        text = text.replace("\r\n", "\n").replace("\r", "\n")
        chunks: list[str] = []
        start = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(text):
                break
            start = end - overlap
        return chunks

    def build_chunks(self, doc_text: str, cfg: RetrievalConfig, grandpa_id: str) -> list["ChildChunk"]:
        parents = self.split_text(doc_text, cfg.parent_chunk_size, cfg.chunk_overlap)
        children: list[ChildChunk] = []
        for parent_text in parents:
            parent_id = _md5_hex(f"{grandpa_id}:{parent_text}")
            child_texts = self.split_text(parent_text, cfg.child_chunk_size, cfg.chunk_overlap)
            for child_text in child_texts:
                child_id = _md5_hex(f"{grandpa_id}:{parent_id}:{child_text}")
                children.append(
                    ChildChunk(
                        id=child_id,
                        grandpa_id=grandpa_id,
                        text=child_text,
                        parent_id=parent_id,
                        parent_context=parent_text,
                    )
                )
        return children


_DOCUMENT_LOADER = DocumentLoader()
_CHUNK_BUILDER = ChunkBuilder()


def _md5_hex(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def _read_text_file(path: Path) -> str:
    return _DOCUMENT_LOADER.read_text_file(path)


def _load_pdf(path: Path) -> str:
    return _DOCUMENT_LOADER.load_pdf(path)


def _load_docx(path: Path) -> str:
    return _DOCUMENT_LOADER.load_docx(path)


def _load_document(path: Path) -> str:
    return _DOCUMENT_LOADER.load(path)


def _split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    return _CHUNK_BUILDER.split_text(text, chunk_size, overlap)


@dataclass(frozen=True)
class ChildChunk:
    id: str
    grandpa_id: str
    text: str
    parent_id: str
    parent_context: str


def _build_chunks(doc_text: str, cfg: RetrievalConfig, grandpa_id: str) -> list[ChildChunk]:
    return _CHUNK_BUILDER.build_chunks(doc_text, cfg, grandpa_id)


def _grandpa_id_for_doc(path: Path, content: str) -> str:
    resolved = str(path.resolve())
    return _md5_hex(f"{resolved}:{_md5_hex(content)}")


def _parse_config(config_path: Path) -> tuple[MilvusConfig, RetrievalConfig]:
    parser = configparser.ConfigParser()
    parser.read(config_path, encoding="utf-8")

    milvus = MilvusConfig(
        host=parser.get("milvus", "host"),
        port=parser.getint("milvus", "port"),
        database_name=parser.get("milvus", "database_name"),
        collection_name=parser.get("milvus", "collection_name"),
    )

    retrieval = RetrievalConfig(
        parent_chunk_size=parser.getint("retrieval", "parent_chunk_size"),
        child_chunk_size=parser.getint("retrieval", "child_chunk_size"),
        chunk_overlap=parser.getint("retrieval", "chunk_overlap"),
        retrieval_k=parser.getint("retrieval", "retrieval_k", fallback=5),
        candidate_m=parser.getint("retrieval", "candidate_m", fallback=2),
    )

    return milvus, retrieval


def _connect_milvus(cfg: MilvusConfig) -> None:
    try:
        connections.connect(alias="default", host=cfg.host, port=str(cfg.port), db_name=cfg.database_name)
    except TypeError:
        connections.connect(alias="default", host=cfg.host, port=str(cfg.port))

    try:
        from pymilvus import db

        existing = set(db.list_database())
        if cfg.database_name and cfg.database_name not in existing:
            db.create_database(cfg.database_name)
        if cfg.database_name:
            db.using_database(cfg.database_name)
    except Exception:
        pass


class MilvusService:
    def __init__(self, cfg: MilvusConfig):
        self.cfg = cfg

    @property
    def collection_name(self) -> str:
        return self.cfg.collection_name

    def ensure_collection(self, dense_dim: int, reset: bool) -> Collection:
        return _ensure_collection(name=self.collection_name, dense_dim=dense_dim, reset=reset)

    def require_collection(self) -> Collection:
        if not utility.has_collection(self.collection_name):
            raise RuntimeError(f"collection 不存在: {self.collection_name}")
        col = Collection(self.collection_name)
        col.load()
        return col

    def delete_by_grandpa_id(self, grandpa_id: str) -> tuple[int, int, int]:
        col = self.require_collection()
        expr = f"grandpa_id == '{grandpa_id}'"
        before = col.query(expr=expr, output_fields=["id"], limit=1)
        col.delete(expr=expr)
        col.flush()
        after = col.query(expr=expr, output_fields=["id"], limit=1)
        return len(before), len(after), col.num_entities

    def hybrid_search(
        self,
        dense_vector: list[float],
        sparse_vector: dict[int, float],
        top_k: int,
        dense_weight: float,
        sparse_weight: float,
    ) -> list[dict[str, Any]]:
        return _hybrid_search(
            milvus_cfg=self.cfg,
            collection_name=self.collection_name,
            dense_vector=dense_vector,
            sparse_vector=sparse_vector,
            top_k=top_k,
            dense_weight=dense_weight,
            sparse_weight=sparse_weight,
        )


def _ensure_collection(
    name: str,
    dense_dim: int,
    reset: bool,
) -> Collection:
    if utility.has_collection(name):
        if reset:
            utility.drop_collection(name)
        else:
            return Collection(name)

    fields = [
        FieldSchema(name="id", dtype=DataType.VARCHAR, is_primary=True, auto_id=False, max_length=64),
        FieldSchema(name="grandpa_id", dtype=DataType.VARCHAR, max_length=64),
        FieldSchema(name="text", dtype=DataType.VARCHAR, max_length=65535),
        FieldSchema(name="dense_vector", dtype=DataType.FLOAT_VECTOR, dim=dense_dim),
        FieldSchema(name="sparse_vector", dtype=DataType.SPARSE_FLOAT_VECTOR),
        FieldSchema(name="parend_id", dtype=DataType.VARCHAR, max_length=64),
        FieldSchema(name="parend_context", dtype=DataType.VARCHAR, max_length=65535),
    ]
    schema = CollectionSchema(fields=fields, description="edurag 初始化集合")
    collection = Collection(name=name, schema=schema)

    try:
        collection.create_index(
            field_name="dense_vector",
            index_params={"index_type": "HNSW", "metric_type": "COSINE", "params": {"M": 16, "efConstruction": 200}},
        )
    except Exception:
        pass

    try:
        collection.create_index(
            field_name="sparse_vector",
            index_params={"index_type": "SPARSE_INVERTED_INDEX", "metric_type": "IP", "params": {}},
        )
    except Exception:
        pass

    return collection


def _batched(items: Sequence[Any], batch_size: int) -> Iterable[Sequence[Any]]:
    if batch_size <= 0:
        raise ValueError("batch_size 必须大于 0")
    for i in range(0, len(items), batch_size):
        yield items[i : i + batch_size]


def _to_milvus_sparse(lexical_weights: dict[Any, Any]) -> dict[str, list[Any]]:
    items: list[tuple[int, float]] = []
    for k, v in lexical_weights.items():
        try:
            idx = int(k)
        except Exception:
            continue
        try:
            weight = float(v)
        except Exception:
            continue
        if weight != 0.0:
            items.append((idx, weight))
    items.sort(key=lambda x: x[0])
    return {"indices": [i for i, _ in items], "values": [w for _, w in items]}


def _lexical_weights_to_csr(lexical_weights: Sequence[dict[Any, Any]], dim: int):
    from scipy.sparse import csr_matrix

    data: list[float] = []
    indices: list[int] = []
    indptr: list[int] = [0]

    for row in lexical_weights:
        row_items: list[tuple[int, float]] = []
        for k, v in row.items():
            try:
                idx = int(k)
                val = float(v)
            except Exception:
                continue
            if val != 0.0:
                row_items.append((idx, val))
        row_items.sort(key=lambda x: x[0])
        for idx, val in row_items:
            indices.append(idx)
            data.append(val)
        indptr.append(len(indices))

    return csr_matrix(
        (np.asarray(data, dtype=np.float32), np.asarray(indices, dtype=np.int64), np.asarray(indptr, dtype=np.int64)),
        shape=(len(lexical_weights), dim),
        dtype=np.float32,
    )


def _infer_sparse_dim(model: Any) -> int:
    for attr_path in ("tokenizer.vocab_size", "model.config.vocab_size", "model.vocab_size", "vocab_size"):
        node: Any = model
        ok = True
        for part in attr_path.split("."):
            if not hasattr(node, part):
                ok = False
                break
            node = getattr(node, part)
        if ok:
            try:
                value = int(node)
                if value > 0:
                    return value
            except Exception:
                pass
    return 250002


def _quiet_embedding_output() -> None:
    os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
    try:
        import warnings

        warnings.filterwarnings("ignore", message=r".*XLMRobertaTokenizerFast tokenizer.*")
    except Exception:
        pass
    try:
        from transformers.utils import logging as hf_logging

        hf_logging.set_verbosity_error()
    except Exception:
        pass


def _bge_encode(model: Any, texts: list[str], **kwargs: Any) -> Any:
    try:
        return model.encode(texts, show_progress_bar=False, **kwargs)
    except TypeError:
        return model.encode(texts, **kwargs)


class BGEEmbedder:
    def __init__(self, model_path: Path):
        self.model_path = model_path
        self._model: Any | None = None
        self._sparse_dim: int | None = None
        _quiet_embedding_output()

    def _get_model(self) -> Any:
        if self._model is not None:
            return self._model
        from FlagEmbedding import BGEM3FlagModel

        use_fp16 = os.environ.get("BGE_USE_FP16", "").strip() == "1"
        self._model = BGEM3FlagModel(str(self.model_path), use_fp16=use_fp16)
        return self._model

    def sparse_dim(self) -> int:
        if self._sparse_dim is not None:
            return self._sparse_dim
        self._sparse_dim = _infer_sparse_dim(self._get_model())
        return self._sparse_dim

    def embed_texts(self, texts: list[str], batch_size: int) -> tuple[list[list[float]], list[dict[Any, Any]], int]:
        model = self._get_model()
        sparse_dim = self.sparse_dim()
        result = _bge_encode(
            model,
            texts,
            batch_size=batch_size,
            max_length=8192,
            return_dense=True,
            return_sparse=True,
            return_colbert_vecs=False,
        )
        dense_vecs_raw = result["dense_vecs"]
        if hasattr(dense_vecs_raw, "tolist"):
            dense_vecs: list[list[float]] = dense_vecs_raw.tolist()
        else:
            dense_vecs = [list(v) for v in dense_vecs_raw]
        lexical_weights = result.get("lexical_weights")
        if lexical_weights is None:
            raise RuntimeError("未生成稀疏向量(lexical_weights)，请检查 FlagEmbedding 版本与参数")

        return dense_vecs, lexical_weights, sparse_dim

    def embed_query(self, query: str) -> tuple[list[float], dict[int, float]]:
        model = self._get_model()
        result = _bge_encode(
            model,
            [query],
            batch_size=1,
            max_length=8192,
            return_dense=True,
            return_sparse=True,
            return_colbert_vecs=False,
        )
        dense_vecs_raw = result["dense_vecs"]
        dense = dense_vecs_raw[0].tolist() if hasattr(dense_vecs_raw[0], "tolist") else list(dense_vecs_raw[0])

        lexical_weights = result.get("lexical_weights")
        if not lexical_weights:
            return dense, {}

        sparse: dict[int, float] = {}
        for k, v in lexical_weights[0].items():
            try:
                idx = int(k)
                val = float(v)
            except Exception:
                continue
            if val != 0.0:
                sparse[idx] = val
        return dense, sparse


_BGE_EMBEDDERS: dict[str, BGEEmbedder] = {}


def _get_bge_embedder(model_path: Path) -> BGEEmbedder:
    key = str(model_path.resolve())
    embedder = _BGE_EMBEDDERS.get(key)
    if embedder is None:
        embedder = BGEEmbedder(model_path)
        _BGE_EMBEDDERS[key] = embedder
    return embedder


def _embed_texts(model_path: Path, texts: list[str], batch_size: int) -> tuple[list[list[float]], list[dict[Any, Any]], int]:
    return _get_bge_embedder(model_path).embed_texts(texts=texts, batch_size=batch_size)


def _embed_query(model_path: Path, query: str) -> tuple[list[float], dict[int, float]]:
    return _get_bge_embedder(model_path).embed_query(query=query)


def _call_llm_from_config(config_path: Path, question: str, context: str, phone: str) -> str:
    parser = configparser.ConfigParser()
    parser.read(config_path, encoding="utf-8")
    provider = parser.get("llm", "provider", fallback="openai_compat").strip() or "openai_compat"
    model = parser.get("llm", "model")
    api_key = (
        parser.get("llm", "api_key", fallback="")
        or parser.get("llm", "openai_api_key", fallback="")
        or parser.get("llm", "dashscope_api_key", fallback="")
        or ""
    ).strip()
    base_url = (
        parser.get("llm", "base_url", fallback="")
        or parser.get("llm", "openai_base_url", fallback="")
        or parser.get("llm", "dashscope_base_url", fallback="")
        or ""
    ).strip()

    if provider == "anthropic":
        import urllib.request
        import json as _json

        url = (parser.get("llm", "anthropic_base_url", fallback="https://api.anthropic.com") + "/v1/messages").strip()
        if not api_key:
            raise RuntimeError("llm 配置不完整：anthropic 需要 anthropic_api_key")
        template = (
            "你是一个智能助手，帮助用户回答问题。\n"
            "优先基于上下文回答；如果上下文不足或与问题不相关，请基于通用知识补充回答，并在答案开头标注【通用知识】。\n"
            "如果答案来源于检索到的文档，请在回答中说明。\n\n"
            f"上下文: {context}\n"
            f"问题: {question}\n\n"
            f"如果无法回答，请回复：“信息不足，无法回答，请联系人工客服，电话：{phone}。”\n"
            "回答:\n"
        )
        data = _json.dumps({"model": model, "max_tokens": 1024, "messages": [{"role": "user", "content": template}]})
        req = urllib.request.Request(
            url=url,
            data=data.encode("utf-8"),
            headers={
                "content-type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=60) as resp:
            body = resp.read().decode("utf-8")
            obj = _json.loads(body)
            try:
                return (obj["content"][0]["text"] or "").strip()
            except Exception:
                return ""

    if not api_key or not base_url:
        raise RuntimeError("llm 配置不完整：需要 api_key 与 base_url")

    from openai import OpenAI
    template = (
        "你是一个智能助手，帮助用户回答问题。\n"
        "优先基于上下文回答；如果上下文不足或与问题不相关，请基于通用知识补充回答，并在答案开头标注【通用知识】。\n"
        "如果答案来源于检索到的文档，请在回答中说明。\n\n"
        f"上下文: {context}\n"
        f"问题: {question}\n\n"
        f"如果无法回答，请回复：“信息不足，无法回答，请联系人工客服，电话：{phone}。”\n"
        "回答:\n"
    )
    prompt = template
    client = OpenAI(api_key=api_key, base_url=base_url)
    resp = client.chat.completions.create(model=model, messages=[{"role": "user", "content": prompt}])
    return (resp.choices[0].message.content or "").strip()


def _hybrid_search(
    milvus_cfg: MilvusConfig,
    collection_name: str,
    dense_vector: list[float],
    sparse_vector: dict[int, float],
    top_k: int,
    dense_weight: float,
    sparse_weight: float,
) -> list[dict[str, Any]]:
    from pymilvus import AnnSearchRequest, MilvusClient, WeightedRanker

    client = MilvusClient(uri=f"http://{milvus_cfg.host}:{milvus_cfg.port}", db_name=milvus_cfg.database_name)

    dense_req = AnnSearchRequest(
        data=[dense_vector],
        anns_field="dense_vector",
        param={"metric_type": "COSINE", "params": {"nprobe": 10}},
        limit=top_k,
    )
    sparse_req = AnnSearchRequest(
        data=[sparse_vector],
        anns_field="sparse_vector",
        param={"metric_type": "IP", "params": {}},
        limit=top_k,
    )
    ranker = WeightedRanker(dense_weight, sparse_weight)
    result = client.hybrid_search(
        collection_name=collection_name,
        reqs=[dense_req, sparse_req],
        ranker=ranker,
        limit=top_k,
        output_fields=["id", "grandpa_id", "text", "parend_id", "parend_context"],
    )
    return result[0]


def _unique_parents(hits: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen: set[str] = set()
    parents: list[dict[str, Any]] = []
    for hit in hits:
        entity = hit.get("entity") or hit
        parent_id = entity.get("parend_id") or ""
        if not parent_id or parent_id in seen:
            continue
        seen.add(parent_id)
        parents.append(entity)
    return parents


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--config", default="config.ini")
    parser.add_argument("--model-path", default=str(Path("model") / "bge-m3"))
    parser.add_argument(
        "--docs",
        nargs="*",
        default=["人工智能就业课课程大纲.md", "人工智能就业课课程大纲.txt"],
    )
    parser.add_argument("--delete-grandpa-id", default=None)
    parser.add_argument("--question", default=None)
    parser.add_argument("--phone", default="10086")
    parser.add_argument("--top-k", type=int, default=20)
    parser.add_argument("--dense-weight", type=float, default=1.0)
    parser.add_argument("--sparse-weight", type=float, default=0.8)
    parser.add_argument("--no-llm", action="store_true")
    parser.add_argument("--show-hits", action="store_true")
    parser.add_argument("--reset", action="store_true")
    parser.add_argument("--batch-size", type=int, default=16)
    return parser.parse_args()


def main() -> int:
    args = _parse_args()
    root = Path(__file__).resolve().parent
    config_path = (root / args.config).resolve()
    model_path = (root / args.model_path).resolve()
    doc_paths = [(root / p).resolve() for p in args.docs]

    if not config_path.exists():
        raise FileNotFoundError(f"找不到配置文件: {config_path}")

    milvus_cfg, retrieval_cfg = _parse_config(config_path)
    _connect_milvus(milvus_cfg)
    milvus = MilvusService(milvus_cfg)

    if args.delete_grandpa_id:
        before, after, total = milvus.delete_by_grandpa_id(args.delete_grandpa_id)
        print(
            f"删除完成: grandpa_id={args.delete_grandpa_id} before_match={before} after_match={after} 当前总数={total}"
        )
        return 0

    if args.question:
        if not model_path.exists():
            raise FileNotFoundError(f"找不到模型路径: {model_path}")
        milvus.require_collection()

        dense_q, sparse_q = _embed_query(model_path, args.question)
        hits = milvus.hybrid_search(
            dense_vector=dense_q,
            sparse_vector=sparse_q,
            top_k=args.top_k,
            dense_weight=args.dense_weight,
            sparse_weight=args.sparse_weight,
        )
        parents = _unique_parents(hits)
        context = "\n\n".join([p.get("parend_context", "") for p in parents if p.get("parend_context")])
        if args.no_llm:
            if args.show_hits:
                for i, hit in enumerate(hits, start=1):
                    entity = hit.get("entity") if isinstance(hit, dict) else getattr(hit, "entity", None)
                    if not entity:
                        entity = getattr(hit, "entity", {}) if hasattr(hit, "entity") else {}
                    distance = hit.get("distance") if isinstance(hit, dict) else getattr(hit, "distance", None)
                    text = entity.get("text", "")
                    print(f"[{i}] distance={distance}\n{text}\n")
            else:
                print(context)
            return 0
        try:
            answer = _call_llm_from_config(config_path, args.question, context, args.phone)
            print(answer)
        except Exception as e:
            _ = e
            print(f"信息不足，无法回答，请联系人工客服，电话：{args.phone}。")
        return 0

    if not model_path.exists():
        raise FileNotFoundError(f"找不到模型路径: {model_path}")
    for p in doc_paths:
        if not p.exists():
            raise FileNotFoundError(f"找不到文档: {p}")

    all_children: list[ChildChunk] = []
    for p in doc_paths:
        doc_text = _load_document(p)
        grandpa_id = _grandpa_id_for_doc(p, doc_text)
        all_children.extend(_build_chunks(doc_text, retrieval_cfg, grandpa_id))

    if not all_children:
        raise RuntimeError("未生成任何子块，无法初始化")

    dense_vecs, lexical_weights, sparse_dim = _embed_texts(model_path, [c.text for c in all_children], args.batch_size)
    dense_dim = len(dense_vecs[0]) if dense_vecs else 0
    if dense_dim <= 0:
        raise RuntimeError("未生成稠密向量，无法初始化")

    collection = milvus.ensure_collection(dense_dim=dense_dim, reset=args.reset)

    ids = [c.id for c in all_children]
    grandpa_ids = [c.grandpa_id for c in all_children]
    texts = [c.text for c in all_children]
    parent_ids = [c.parent_id for c in all_children]
    parent_contexts = [c.parent_context for c in all_children]

    for batch in _batched(list(range(len(all_children))), args.batch_size):
        b_ids = [ids[i] for i in batch]
        b_grandpa_ids = [grandpa_ids[i] for i in batch]
        b_texts = [texts[i] for i in batch]
        b_dense = [dense_vecs[i] for i in batch]
        b_lexical = [lexical_weights[i] for i in batch]
        b_sparse = _lexical_weights_to_csr(b_lexical, sparse_dim)
        b_parent_ids = [parent_ids[i] for i in batch]
        b_parent_contexts = [parent_contexts[i] for i in batch]
        collection.insert([b_ids, b_grandpa_ids, b_texts, b_dense, b_sparse, b_parent_ids, b_parent_contexts])

    collection.flush()
    collection.load()
    total = collection.num_entities

    print(
        f"初始化完成: collection={milvus_cfg.collection_name} dense_dim={dense_dim} sparse_vocab_hint={sparse_dim} 写入子块={len(all_children)} 当前总数={total}"
    )

    if total <= 0:
        raise RuntimeError("写入后集合仍为空，初始化失败")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

